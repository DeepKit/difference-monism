import re
from pathlib import Path
from typing import Optional, Union
import mimetypes


class TextExtractor:
    """Extract text from various file formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.txt': self._extract_txt,
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.html': self._extract_html,
            '.htm': self._extract_html,
            '.md': self._extract_txt,
            '.csv': self._extract_txt,
            '.json': self._extract_txt,
            '.xml': self._extract_txt,
        }
    
    def extract(self, file_path: Union[str, Path]) -> str:
        """Extract text from file"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {suffix}")
        
        extractor = self.supported_formats[suffix]
        return extractor(path)
    
    def extract_from_string(self, content: str, format_type: str = 'html') -> str:
        """Extract text from string content"""
        if format_type == 'html':
            return self._clean_html(content)
        return content
    
    def _extract_txt(self, path: Path) -> str:
        """Extract from plain text files"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Unable to decode file: {path}")
    
    def _extract_pdf(self, path: Path) -> str:
        """Extract from PDF files"""
        try:
            import PyPDF2
            
            text = []
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            
            return '\n'.join(text)
        except ImportError:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
    
    def _extract_docx(self, path: Path) -> str:
        """Extract from Word documents"""
        try:
            from docx import Document
            
            doc = Document(path)
            text = []
            
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            return '\n'.join(text)
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    def _extract_html(self, path: Path) -> str:
        """Extract from HTML files"""
        content = self._extract_txt(path)
        return self._clean_html(content)
    
    def _clean_html(self, html_content: str) -> str:
        """Remove HTML tags and clean text"""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for script in soup(['script', 'style']):
                script.decompose()
            
            text = soup.get_text()
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
        except ImportError:
            text = re.sub(r'<[^>]+>', '', html_content)
            text = re.sub(r'\s+', ' ', text)
            return text.strip()
    
    def extract_with_metadata(self, file_path: Union[str, Path]) -> dict:
        """Extract text with file metadata"""
        path = Path(file_path)
        
        return {
            'text': self.extract(path),
            'filename': path.name,
            'size': path.stat().st_size,
            'format': path.suffix,
            'mime_type': mimetypes.guess_type(path)[0]
        }


# 使用示例
if __name__ == '__main__':
    extractor = TextExtractor()
    
    # 从文件提取
    # text = extractor.extract('document.pdf')
    # print(text)
    
    # 从HTML字符串提取
    html = '<html><body><h1>标题</h1><p>内容</p></body></html>'
    text = extractor.extract_from_string(html, 'html')
    print(text)
    
    # 提取带元数据
    # result = extractor.extract_with_metadata('document.docx')
    # print(result)